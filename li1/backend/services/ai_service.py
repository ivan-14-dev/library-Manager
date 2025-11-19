# services/ai_service.py
import openai
from django.conf import settings
from django.utils import timezone
from ..models import AIConfiguration
import logging

logger = logging.getLogger(__name__)

class AIService:
    """
    Service pour l'intégration avec les APIs d'IA (OpenAI, etc.)
    Gère les limites d'usage et la configuration par utilisateur
    """
    
    def __init__(self):
        self.openai_client = openai.OpenAI(api_key=settings.OPENAI_API_KEY)
    
    def check_usage_limits(self, user, feature):
        """
        Vérifier si l'utilisateur peut utiliser une fonctionnalité IA
        Retourne (can_use: bool, reason: str)
        """
        try:
            config = AIConfiguration.objects.get(user=user)
            
            # Vérifier si l'IA est activée
            if not config.is_active:
                return False, "L'IA n'est pas activée pour votre compte"
            
            # Vérifier si la fonctionnalité est autorisée
            if feature not in config.allowed_features:
                return False, f"La fonctionnalité '{feature}' n'est pas incluse dans votre abonnement"
            
            # Vérifier les limites d'usage
            current_usage = config.current_usage.get(feature, 0)
            usage_limit = config.usage_limits.get(feature, 0)
            
            if usage_limit > 0 and current_usage >= usage_limit:
                return False, f"Limite d'usage atteinte pour '{feature}'"
            
            return True, "OK"
            
        except AIConfiguration.DoesNotExist:
            return False, "Aucune configuration IA trouvée"
    
    def increment_usage(self, user, feature):
        """Incrémenter le compteur d'usage pour une fonctionnalité"""
        try:
            config = AIConfiguration.objects.get(user=user)
            config.current_usage[feature] = config.current_usage.get(feature, 0) + 1
            config.save()
        except AIConfiguration.DoesNotExist:
            logger.warning(f"Tentative d'incrémenter l'usage pour un utilisateur sans config: {user.username}")
    
    def process_grammar_check(self, user, text):
        """
        Vérification grammaticale avec IA
        Retourne les corrections suggérées
        """
        # Vérifier les limites
        can_use, reason = self.check_usage_limits(user, 'grammar_check')
        if not can_use:
            raise PermissionError(reason)
        
        try:
            response = self.openai_client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "Vous êtes un assistant de correction grammaticale. Corrigez les erreurs dans le texte fourni et expliquez brièvement les corrections."},
                    {"role": "user", "content": f"Corrigez ce texte : {text}"}
                ],
                max_tokens=500
            )
            
            # Incrémenter l'usage
            self.increment_usage(user, 'grammar_check')
            
            return {
                'corrected_text': response.choices[0].message.content,
                'original_text': text,
                'model_used': 'gpt-3.5-turbo'
            }
            
        except Exception as e:
            logger.error(f"Erreur lors de la vérification grammaticale: {str(e)}")
            raise
    
    def generate_content(self, user, prompt, context=None):
        """
        Génération de contenu avec IA
        """
        can_use, reason = self.check_usage_limits(user, 'content_generation')
        if not can_use:
            raise PermissionError(reason)
        
        try:
            system_message = "Vous êtes un assistant d'écriture créative. Aidez à générer du contenu basé sur le prompt."
            if context:
                system_message += f" Contexte: {context}"
            
            response = self.openai_client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": system_message},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=1000
            )
            
            self.increment_usage(user, 'content_generation')
            
            return {
                'generated_content': response.choices[0].message.content,
                'prompt': prompt,
                'model_used': 'gpt-3.5-turbo'
            }
            
        except Exception as e:
            logger.error(f"Erreur lors de la génération de contenu: {str(e)}")
            raise
    
    def analyze_sentiment(self, user, text):
        """
        Analyse de sentiment avec IA
        """
        can_use, reason = self.check_usage_limits(user, 'sentiment_analysis')
        if not can_use:
            raise PermissionError(reason)
        
        try:
            response = self.openai_client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "Analysez le sentiment du texte et fournissez un score entre -1 (négatif) et 1 (positif), ainsi qu'une brève analyse."},
                    {"role": "user", "content": text}
                ],
                max_tokens=300
            )
            
            self.increment_usage(user, 'sentiment_analysis')
            
            return {
                'analysis': response.choices[0].message.content,
                'text_analyzed': text,
                'model_used': 'gpt-3.5-turbo'
            }
            
        except Exception as e:
            logger.error(f"Erreur lors de l'analyse de sentiment: {str(e)}")
            raise

# services/export_service.py
from weasyprint import HTML
from docx import Document
from celery import shared_task
from django.conf import settings
import os
import tempfile
from ..models import ExportJob

class ExportService:
    """
    Service pour la génération de fichiers d'export (PDF, DOCX, etc.)
    """
    
    def __init__(self):
        self.export_dir = getattr(settings, 'EXPORT_ROOT', os.path.join(settings.MEDIA_ROOT, 'exports'))
        os.makedirs(self.export_dir, exist_ok=True)
    
    def generate_pdf(self, content, options=None):
        """
        Générer un PDF à partir du contenu HTML
        """
        options = options or {}
        
        try:
            # Créer un HTML complet avec styles
            html_content = self._wrap_content_for_export(content, options, 'pdf')
            
            # Générer le PDF
            html = HTML(string=html_content)
            pdf_file = html.write_pdf()
            
            return pdf_file
            
        except Exception as e:
            logger.error(f"Erreur lors de la génération PDF: {str(e)}")
            raise
    
    def generate_docx(self, content, options=None):
        """
        Générer un document Word à partir du contenu
        """
        options = options or {}
        
        try:
            doc = Document()
            
            # Convertir le HTML en structure Word
            # Note: Cette conversion est basique, vous pouvez l'améliorer
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(content, 'html.parser')
            
            for element in soup.find_all(['p', 'h1', 'h2', 'h3']):
                if element.name == 'p':
                    doc.add_paragraph(element.get_text())
                elif element.name.startswith('h'):
                    doc.add_heading(element.get_text(), level=int(element.name[1]))
            
            # Sauvegarder dans un fichier temporaire
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                doc.save(tmp_file.name)
                with open(tmp_file.name, 'rb') as f:
                    docx_content = f.read()
                
                os.unlink(tmp_file.name)
            
            return docx_content
            
        except Exception as e:
            logger.error(f"Erreur lors de la génération DOCX: {str(e)}")
            raise
    
    def generate_html(self, content, options=None):
        """
        Générer un fichier HTML complet
        """
        options = options or {}
        return self._wrap_content_for_export(content, options, 'html').encode('utf-8')
    
    def generate_markdown(self, content, options=None):
        """
        Convertir le HTML en Markdown
        """
        options = options or {}
        
        try:
            from markdownify import markdownify as md
            markdown_content = md(content)
            return markdown_content.encode('utf-8')
            
        except Exception as e:
            logger.error(f"Erreur lors de la conversion Markdown: {str(e)}")
            raise
    
    def _wrap_content_for_export(self, content, options, format_type):
        """
        Encapsuler le contenu dans un template HTML complet
        """
        title = options.get('title', 'Document Exporté')
        include_header = options.get('include_header', True)
        include_footer = options.get('include_footer', True)
        
        header = ""
        if include_header:
            header = f"""
            <div class="header">
                <h1>{title}</h1>
                <div class="meta">
                    Exporté le {timezone.now().strftime('%d/%m/%Y à %H:%M')} • Bibliothèque Numérique
                </div>
                <hr>
            </div>
            """
        
        footer = ""
        if include_footer:
            footer = f"""
            <div class="footer">
                <hr>
                <div class="footer-meta">
                    Page <span class="page-number"></span> • Document généré par Bibliothèque Numérique
                </div>
            </div>
            """
        
        css = self._get_css_for_format(format_type)
        
        return f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>{title}</title>
            <style>{css}</style>
        </head>
        <body>
            {header}
            <div class="content">
                {content}
            </div>
            {footer}
        </body>
        </html>
        """
    
    def _get_css_for_format(self, format_type):
        """CSS adapté selon le format d'export"""
        if format_type == 'pdf':
            return """
            body {{ font-family: Arial, sans-serif; line-height: 1.6; margin: 2cm; }}
            .header {{ text-align: center; margin-bottom: 2cm; }}
            .footer {{ margin-top: 2cm; text-align: center; font-size: 0.8em; color: #666; }}
            .content {{ margin: 1cm 0; }}
            h1, h2, h3 {{ color: #333; }}
            table {{ border-collapse: collapse; width: 100%; }}
            table, th, td {{ border: 1px solid #ddd; }}
            th, td {{ padding: 8px; text-align: left; }}
            """
        else:
            return """
            body {{ font-family: Arial, sans-serif; line-height: 1.6; max-width: 800px; margin: 0 auto; padding: 20px; }}
            .header {{ text-align: center; margin-bottom: 2em; }}
            .footer {{ margin-top: 2em; text-align: center; font-size: 0.9em; color: #666; }}
            .content {{ margin: 1em 0; }}
            h1, h2, h3 {{ color: #333; }}
            table {{ border-collapse: collapse; width: 100%; }}
            table, th, td {{ border: 1px solid #ddd; }}
            th, td {{ padding: 8px; text-align: left; }}
            """

@shared_task
def process_export_job(job_id):
    """
    Tâche Celery pour traiter un job d'exportation asynchrone
    """
    try:
        job = ExportJob.objects.get(id=job_id)
        job.status = 'processing'
        job.save()
        
        export_service = ExportService()
        
        # Générer le fichier selon le format
        if job.export_format == 'pdf':
            file_content = export_service.generate_pdf(job.document.content, job.options)
            file_extension = '.pdf'
        elif job.export_format == 'docx':
            file_content = export_service.generate_docx(job.document.content, job.options)
            file_extension = '.docx'
        elif job.export_format == 'html':
            file_content = export_service.generate_html(job.document.content, job.options)
            file_extension = '.html'
        elif job.export_format == 'markdown':
            file_content = export_service.generate_markdown(job.document.content, job.options)
            file_extension = '.md'
        else:
            raise ValueError(f"Format non supporté: {job.export_format}")
        
        # Sauvegarder le fichier
        filename = f"export_{job.id}_{int(timezone.now().timestamp())}{file_extension}"
        filepath = os.path.join(export_service.export_dir, filename)
        
        with open(filepath, 'wb') as f:
            f.write(file_content)
        
        # Mettre à jour le job
        job.status = 'completed'
        job.file_path = filepath
        job.progress = 100
        job.completed_at = timezone.now()
        job.save()
        
    except Exception as e:
        logger.error(f"Erreur lors du traitement du job d'export {job_id}: {str(e)}")
        
        job.status = 'failed'
        job.error_message = str(e)
        job.save()